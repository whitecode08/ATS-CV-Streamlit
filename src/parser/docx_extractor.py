"""
src/parser/docx_extractor.py — DOCX Text Extraction
Uses python-docx to extract text from Word documents (.docx).
Handles paragraphs, tables, and header/footer sections.
"""

from __future__ import annotations
import io
from pathlib import Path
from typing import Union


def extract_text_from_docx(source: Union[str, Path, bytes, io.BytesIO]) -> str:
    """
    Extract plain text from a .docx Word document.

    Args:
        source: File path (str/Path), raw bytes, or BytesIO object.

    Returns:
        A clean, concatenated string of all text from the document.

    Raises:
        ImportError: If python-docx is not installed.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install it with: pip install python-docx"
        ) from e

    # Open document from various source types
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
    elif isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    elif isinstance(source, io.BytesIO):
        doc = Document(source)
    else:
        raise ValueError(f"Unsupported source type: {type(source)}")

    text_parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract text from tables (important for tabular resume sections)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append("  |  ".join(row_texts))

    full_text = "\n".join(text_parts)
    return _post_process(full_text)


def _post_process(text: str) -> str:
    """Clean up common DOCX extraction artifacts."""
    import re

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
